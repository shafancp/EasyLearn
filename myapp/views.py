import datetime

import PyPDF2
import docx
import requests
# import nltk
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render
import markdown
from docx import Document
from io import BytesIO
from fpdf import FPDF
from django.utils.html import strip_tags

# Create your views here.
from myapp.models import *


def myapp(request):
    return HttpResponse('hello world')


def login(request):
    return render(request,"login.html")

def login_post(request):
    username=request.POST["textfield"]
    password=request.POST["textfield2"]
    SD=Login.objects.filter(username=username,password=password)
    if SD.exists():
        DS=Login.objects.get(username=username,password=password)
        request.session['lid']=DS.id
        if DS.type=='admin':
            return HttpResponse('''<script>alert("successfully login");window.location='/myapp/adminhome/'</script>''')
        elif DS.type=='student':
            return HttpResponse('''<script>alert("successfully login");window.location='/myapp/studenthome/'</script>''')
        elif DS.type=='tutor':
            return HttpResponse(
                '''<script>alert("successfully login");window.location='/myapp/tutor_home/'</script>''')
        else:
            return HttpResponse('''<script>alert("invalid username and password");window.location='/myapp/login/'</script>''')
    else:
        return HttpResponse(
            '''<script>alert("not found username and password");window.location='/myapp/login/'</script>''')


def adminhome (request):
    return render(request, "admin/admin_home.html")


def addsubject(request):
    return render(request, "admin/addsubject.html")

def change_password(request):
    return render(request, "admin/changepassword.html")
def change_password_post(request):
    oldpassword = request.POST["textfield"]
    newpassword = request.POST["textfield2"]
    confirmpassword = request.POST["textfield3"]
    obj =  Login.objects.filter(id = request.session['lid'],password=oldpassword)
    if obj.exists():
        if newpassword == confirmpassword:
            Login.objects.filter(id=request.session['lid']).update(password=newpassword)
            return HttpResponse('''<script>alert("Password Changed");window.location='/myapp/change_password/'</script>''')
        else:
            return HttpResponse('''<script>alert("Password must be same");window.location='/myapp/change_password/'</script>''')
    else:
        return HttpResponse('''<script>alert("Old Password is incorrect");window.location='/myapp/change_password/'</script>''')


def addsubject_post(request):
    subject = request.POST["textfield"]
    obj=Subject()
    obj.name=subject
    obj.save()
    return HttpResponse('''<script>alert("subject added");window.location='/myapp/addsubject/'</script>''')


def delete_subject(request,id):
    res=Subject.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("Delete Successfull");window.location='/myapp/viewsubject/'</script>''')



def approvedtutors(request):
    AS=Tutor.objects.filter(status='approved')
    return render(request, "admin/approvedtutors.html",{'data':AS})


def approvedtutors_post(request):
    search=request.POST["textfield"]
    AS = Tutor.objects.filter(status='approved',firstname__icontains=search)
    return render(request, "admin/approvedtutors.html", {'data': AS})


def rejectedtutors(request):
    AS=Tutor.objects.filter(status='rejected')
    return render(request, "admin/rejectedtutors.html",{'data':AS})


def rejectedtutors_post(request):
    search=request.POST["textfield"]
    AS = Tutor.objects.filter(status='rejected',firstname__icontains=search)
    return render(request, "admin/rejectedtutors.html", {'data': AS})


def viewcomplaint(request):
    AS = Complaint.objects.all()
    return render(request, "admin/viewcomplaint.html",{'data':AS})


def viewcomplaint_post(request):
    fromdate=request.POST["textfield"]
    todate=request.POST["textfield2"]
    AS = Complaint.objects.filter(date__range=[fromdate,todate])
    return render(request, "admin/viewcomplaint.html", {'data': AS})

def viewfeedback(request):
    AS = Feedback.objects.all()
    return render(request, "admin/viewfeedback.html",{'data':AS})


def viewfeedback_post(request):
    fromdate = request.POST["textfield"]
    todate = request.POST["textfield2"]
    AS = Feedback.objects.filter(date__range=[fromdate, todate])
    return render(request, "admin/viewfeedback.html", {'data': AS})

def viewnotes_admin(request):
    res=Subject.objects.all()
    AS = Tutor.objects.all()
    de=Notes.objects.all()
    return render(request, "admin/viewnotesadmin.html", {'data':de, 'data1':res, 'data2':AS})

def viewnotes_admin_post(request):
    subject = request.POST["select"]
    tutor = request.POST["select2"]
    search = request.POST["textfield"]
    res=Subject.objects.all()
    AS = Tutor.objects.all()
    de=Notes.objects.all()
    if subject == "" and tutor=="" and search=="":
        de = Notes.objects.all()
    elif tutor == "" and search=="":
        de = Notes.objects.filter(TUTOR__SUBJECT__id=subject)
    elif tutor=="" and subject =="":
        de = Notes.objects.filter(topic__icontains=search)
    elif subject=="" and search=="":
        de = Notes.objects.filter(TUTOR__id=tutor)
    elif search=="":
        de = Notes.objects.filter(TUTOR__id=tutor, TUTOR__SUBJECT__id=subject)
    elif subject=="":
        de = Notes.objects.filter(TUTOR__id=tutor, topic__icontains=search)
    elif tutor=="":
        de = Notes.objects.filter(TUTOR__SUBJECT__id=subject, topic__icontains=search)
    else:
        de = Notes.objects.filter(topic__icontains=search, TUTOR__SUBJECT__id=subject, TUTOR__id=tutor)
    return render(request, "admin/viewnotesadmin.html", {'data': de, 'data1':res, 'data2':AS})

def viewstudent(request):
    do=Student.objects.all()
    return render(request, "admin/viewstudent.html",{'data':do})


def viewstudent_post(request):
    search = request.POST["textfield"]
    AS = Student.objects.filter(firstname__icontains=search)
    return render(request, "admin/viewstudent.html", {'data': AS})

def viewsubject(request):
    ras=Subject.objects.all()
    return render(request, "admin/viewsubject.html",{'data':ras})


def viewsubject_post(request):
    subject=request.POST["textfield"]
    AS = Subject.objects.filter(name__icontains=subject)
    return render(request, "admin/viewsubject.html", {'data': AS})

def viewtutorverify(request):
    fe=Tutor.objects.filter(status="pending")
    return render(request, "admin/viewtutorverify.html",{'data':fe})

def approvetutor(request,id):
    DF=Tutor.objects.filter(id=id).update(status="approved")
    Login.objects.filter(id=Tutor.objects.get(id=id).LOGIN_id).update(type="tutor")
    return HttpResponse('''<script>alert("approved");window.location='/myapp/viewtutorverify/'</script>''')

def rejecttutor(request,id):
    DF=Tutor.objects.filter(id=id).update(status="rejected")
    return HttpResponse('''<script>alert("rejected");window.location='/myapp/viewtutorverify/'</script>''')

def viewtutorverify_post(request):
    search=request.POST["textfield"]
    AS = Tutor.objects.filter(status='pending', firstname__icontains=search)
    return render(request, "admin/viewtutorverify.html", {'data': AS})

def adminreply(request,id):
    return render(request, "admin/adminreply.html",{'data':id})

def adminreply_post(request):
   reply=request.POST["textfield2"]
   id=request.POST["cid"]
   cc=Complaint.objects.filter(id=id).update(reply=reply,status="replied")
   return HttpResponse('''<script>alert("replied");window.location='/myapp/viewcomplaint/'</script>''')




#student
def studenthome(request):
    return render(request, "student/homenew.html")


def registration(request):
    return render(request, "student/studentsignup.html")
def registration_post(request):
    firstname = request.POST["textfield"]
    lastname = request.POST["textfield2"]
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d-%H%M%S')+".jpg"
    photo = request.FILES["fileField"]
    fs = FileSystemStorage()
    fn = fs.save(date,photo)
    path = fs.url(date)
    DOB = request.POST["textfield3"]
    gender = request.POST["RadioGroup1"]
    course = request.POST["textfield4"]
    institution = request.POST["textfield5"]
    email = request.POST["textfield6"]
    phoneno = request.POST["textfield7"]
    password = request.POST["textfield8"]
    confirmpassword = request.POST["confirmpassword"]

    if password != confirmpassword:
        return HttpResponse('''<script>alert("Password do not match");history.back()</script>''')
    if Login.objects.filter(username = email).exists():
        return HttpResponse('''<script>alert("Mail already exists");history.back()</script>''')
    lobj=Login()
    lobj.username=email
    lobj.password=password
    lobj.type='student'
    lobj.save()

    obj = Student()
    obj.firstname=firstname
    obj.lastname=lastname
    obj.photo=path
    obj.DOB=DOB
    obj.gender=gender
    obj.course=course
    obj.institution=institution
    obj.email=email
    obj.phoneno=phoneno
    obj.LOGIN=lobj
    obj.save()
    return HttpResponse('''<script>alert("Registered");window.location='/myapp/login/'</script>''')


def senddoubts(request,id):
    request.session["tid"]=id
    return render(request, "student/senddoubts.html",{'id':id})
def senddoubts_post(request):
    doubt = request.POST['textfield']
    note=request.POST['nid']
    obj = Doubts()
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.doubt=doubt
    obj.reply="pending"
    obj.status="pending"
    obj.date=date
    obj.NOTE_id= note
    obj.STUDENT=Student.objects.get(LOGIN_id=request.session["lid"])
    obj.save()
    return HttpResponse('''<script>alert("doubt sended");window.location='/myapp/viewnotes/'</script>''')

def viewdoubts(request):
    de = Doubts.objects.filter(STUDENT__LOGIN_id=request.session['lid'])
    return render(request, "student/viewdoubts.html",{'data':de})
def viewdoubts_post(request):
    return HttpResponse("viewed")

def viewnotes(request):
    res = Subject.objects.all()
    AS = Tutor.objects.all()
    de = Notes.objects.all()
    return render(request, "student/viewnotes.html", {'data': de, 'data1': res, 'data2': AS})

def viewnotes_post(request):
    subject = request.POST["select"]
    tutor = request.POST["select2"]
    search = request.POST["textfield"]
    res=Subject.objects.all()
    AS = Tutor.objects.all()
    de=Notes.objects.all()
    if subject == "" and tutor=="" and search=="":
        de = Notes.objects.all()
    elif tutor == "" and search=="":
        de = Notes.objects.filter(TUTOR__SUBJECT__id=subject)
    elif tutor=="" and subject =="":
        de = Notes.objects.filter(topic__icontains=search)
    elif subject=="" and search=="":
        de = Notes.objects.filter(TUTOR__id=tutor)
    elif search=="":
        de = Notes.objects.filter(TUTOR__id=tutor, TUTOR__SUBJECT__id=subject)
    elif subject=="":
        de = Notes.objects.filter(TUTOR__id=tutor, topic__icontains=search)
    elif tutor=="":
        de = Notes.objects.filter(TUTOR__SUBJECT__id=subject, topic__icontains=search)
    else:
        de = Notes.objects.filter(topic__icontains=search, TUTOR__SUBJECT__id=subject, TUTOR__id=tutor)
    return render(request, "student/viewnotes.html", {'data': de, 'data1':res, 'data2':AS})


def viewprofile(request):
    res=Student.objects.get(LOGIN=request.session['lid'])
    return render(request, "student/viewprofile.html",{'data':res})

def updateprofile(request):
    res=Student.objects.get(LOGIN=request.session['lid'])
    return render(request, "student/updateprofile.html",{'data':res})

def updateprofile_post(request):
    firstname = request.POST["textfield"]
    lastname = request.POST["textfield2"]
    DOB = request.POST["textfield3"]
    gender = request.POST["RadioGroup1"]
    course = request.POST["textfield4"]
    institution = request.POST["textfield5"]
    email = request.POST["textfield6"]
    phoneno = request.POST["textfield7"]
    if 'fileField' in request.FILES:
        photo = request.FILES["fileField"]
        if photo !="":
            from datetime import datetime
            date = datetime.now().strftime('%Y%m%d-%H%M%S') + ".jpg"
            fs = FileSystemStorage()
            fn = fs.save(date, photo)
            path = fs.url(date)


            obj = Student.objects.filter(LOGIN=request.session['lid']).update(firstname=firstname,
                                                                              lastname=lastname,
                                                                              photo=path,
                                                                              DOB=DOB,
                                                                              gender=gender,
                                                                              course=course,
                                                                              institution=institution,
                                                                              email=email,
                                                                              phoneno=phoneno)
    else:
        obj = Student.objects.filter(LOGIN=request.session['lid']).update(firstname=firstname,
                                                                          lastname=lastname,
                                                                          DOB=DOB,
                                                                          gender=gender,
                                                                          course=course,
                                                                          institution=institution,
                                                                          email=email,
                                                                          phoneno=phoneno)



    return HttpResponse('''<script>alert("updated");window.location='/myapp/viewprofile/'</script>''')

def viewtutor(request,id):
    de = Tutor.objects.get(id=id)
    return render(request, "student/viewtutor.html",{'data':de})

def sendfeedback_student(request):
    return render(request, "student/sendfeedback_student.html")
def sendfeedback_student_post(request):
    feedback = request.POST["textfield"]
    obj = Feedback()
    obj.feedback = feedback
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.date = date
    obj.FROM_id = request.session['lid']
    obj.type="student"
    obj.save()
    return HttpResponse('''<script>alert("Feedback Sended");window.location='/myapp/sendfeedback/'</script>''')




def student_sendcomplaint(request):
    return render(request, "student/sendcomplaint.html")
def student_sendcomplaint_post(request):
    head = request.POST["textfield"]
    complaint = request.POST["textfield2"]
    obj = Complaint()
    obj.head = head
    obj.complaint = complaint
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.reply = "pending"
    obj.status = "pending"
    obj.date = date
    obj.FROM_id = request.session['lid']
    obj.save()
    return HttpResponse('''<script>alert("Complaint Sended");window.location='/myapp/student_sendcomplaint/'</script>''')


def student_sendfeedback(request):
    return render(request, "student/sendfeedback.html")
def student_sendfeedback_post(request):
    feedback = request.POST["textfield"]
    obj = Feedback()
    obj.feedback = feedback
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.date = date
    obj.FROM_id = request.session['lid']
    obj.type="student"
    obj.save()
    return HttpResponse('''<script>alert("Feedback Sended");window.location='/myapp/student_sendfeedback/'</script>''')

def student_viewreply(request):
    de=Complaint.objects.filter(FROM_id=request.session['lid'])
    return render(request, "student/viewreply.html",{'data':de})

#tutor
def addnote(request):
    return render(request, "tutor/addnote.html")
def addnote_post(request):
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d-%H%M%S') + ".pdf"
    note = request.FILES["filefield"]
    fs = FileSystemStorage()
    fn = fs.save(date, note)
    path = fs.url(date)
    topic = request.POST["textfield2"]
    description = request.POST["textfield3"]

    obj = Notes()
    date = datetime.now().strftime('%Y%m%d')
    obj.note = path
    obj.topic = topic
    obj.description = description
    obj.date = date
    id = request.session['lid']
    de = Tutor.objects.get(LOGIN=id)
    obj.TUTOR = de
    obj.save()
    return HttpResponse('''<script>alert("Note Added");window.location='/myapp/addnote/'</script>''')

def tutor_home(request):
    return render(request, "tutor/index.html")


def tutor_change_password(request):
    return render(request, "tutor/changepassword.html")
def tutor_change_password_post(request):
    oldpassword = request.POST["textfield"]
    newpassword = request.POST["textfield2"]
    confirmpassword = request.POST["textfield3"]
    obj =  Login.objects.filter(id = request.session['lid'],password=oldpassword)
    if obj.exists():
        if newpassword == confirmpassword:
            Login.objects.filter(id=request.session['lid']).update(password=newpassword)
            return HttpResponse('''<script>alert("Password Changed");window.location='/myapp/change_password/'</script>''')
        else:
            return HttpResponse('''<script>alert("Password must be same");window.location='/myapp/change_password/'</script>''')
    else:
        return HttpResponse('''<script>alert("Old Password is incorrect");window.location='/myapp/change_password/'</script>''')

def registration_tutor(request):
    ds = Subject.objects.all()
    return render(request, "tutor/tutorsignup.html",{'data':ds})
def registration_tutor_post(request):
    firstname = request.POST["textfield"]
    lastname = request.POST["textfield2"]
    from datetime import datetime
    date = 'photo/'+datetime.now().strftime('%Y%m%d-%H%M%S')+".jpg"
    photo = request.FILES["fileField2"]
    fs = FileSystemStorage()
    fn = fs.save(date,photo)
    path = fs.url(date)
    idproof = request.FILES["fileField"]
    date1 = 'id/'+datetime.now().strftime('%Y%m%d-%H%M%S')+".jpg"
    fs1 = FileSystemStorage()
    fn1 = fs1.save(date1,idproof)
    path1 = fs1.url(date1)
    DOB = request.POST["textfield13"]
    gender = request.POST["RadioGroup1"]
    email = request.POST["textfield4"]
    phoneno = request.POST["textfield5"]
    house = request.POST["textfield6"]
    place = request.POST["textfield7"]
    district = request.POST["textfield8"]
    state = request.POST["textfield9"]
    pin = request.POST["textfield10"]
    country = request.POST["textfield11"]
    qualification = request.POST["textfield12"]
    subject = request.POST["select"]
    ss = Subject.objects.get(id=subject)
    password = request.POST["password"]
    confirmpassword = request.POST["confirmpassword"]

    if password != confirmpassword:
        return HttpResponse('''<script>alert("Password do not match");history.back()</script>''')
    if Login.objects.filter(username = email).exists():
        return HttpResponse('''<script>alert("Mail already exists");history.back()</script>''')
    lobj = Login()
    lobj.username = email
    lobj.password = password
    lobj.type = 'pending'
    lobj.save()

    obj = Tutor()
    obj.firstname = firstname
    obj.lastname = lastname
    obj.photo = path
    obj.idproof = path1
    obj.DOB = DOB
    obj.gender = gender
    obj.email = email
    obj.phoneno = phoneno
    obj.houseno = house
    obj.place = place
    obj.district = district
    obj.state = state
    obj.PIN = pin
    obj.country = country
    obj.qualification = qualification
    obj.SUBJECT = ss
    obj.status = 'pending'
    obj.LOGIN = lobj
    obj.save()

    return HttpResponse('''<script>alert("Registered");window.location='/myapp/login/'</script>''')

def sendfeedback(request):
    return render(request, "tutor/sendfeedback.html")
def sendfeedback_post(request):
    feedback = request.POST["textfield"]
    obj = Feedback()
    obj.feedback = feedback
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.date = date
    obj.FROM_id = request.session['lid']
    obj.type="tutor"
    obj.save()
    return HttpResponse('''<script>alert("Feedback Sended");window.location='/myapp/sendfeedback/'</script>''')


def sendcomplaint(request):
    return render(request, "tutor/sendcomplaint.html")
def sendcomplaint_post(request):
    head = request.POST["textfield"]
    complaint = request.POST["textfield2"]
    obj = Complaint()
    obj.head = head
    obj.complaint = complaint
    from datetime import datetime
    date = datetime.now().strftime('%Y%m%d')
    obj.reply = "pending"
    obj.status = "pending"
    obj.date = date
    obj.FROM_id = request.session['lid']
    obj.save()
    return HttpResponse('''<script>alert("Complaint Sended");window.location='/myapp/sendcomplaint/'</script>''')


def viewprofile_tutor(request):
    id = request.session['lid']
    de = Tutor.objects.get(LOGIN=id)
    return render(request, "tutor/viewprofile.html", {'data': de})

def updateprofile_tutor(request):
    id = request.session['lid']
    res = Tutor.objects.get(LOGIN=id)
    re=Subject.objects.all()
    return render(request, "tutor/updateprofile.html", {'data': res,'data2':re})

def updateprofile_tutor_post(request):
    firstname = request.POST["textfield"]
    lastname = request.POST["textfield2"]
    DOB = request.POST["textfield13"]
    gender = request.POST["RadioGroup1"]
    email = request.POST["textfield4"]
    phoneno = request.POST["textfield5"]
    house = request.POST["textfield6"]
    place = request.POST["textfield7"]
    district = request.POST["textfield8"]
    state = request.POST["textfield9"]
    pin = request.POST["textfield10"]
    country = request.POST["textfield11"]
    qualification = request.POST["textfield12"]
    subject = request.POST["select"]
    ss = Subject.objects.get(id=subject)
    if 'fileField' in request.FILES:
        photo = request.FILES["fileField2"]
        if photo !="":
            from datetime import datetime
            date = datetime.now().strftime('%Y%m%d-%H%M%S') + ".jpg"
            fs = FileSystemStorage()
            fn = fs.save(date, photo)
            path = fs.url(date)
            obj = Tutor.objects.filter(LOGIN=request.session['lid']).update(firstname=firstname,lastname=lastname,photo=path,DOB=DOB,gender=gender,
                                                                    email=email,phoneno=phoneno,houseno=house,place=place,district=district,state=state,PIN=pin,
                                                                    country = country,qualification=qualification,SUBJECT=ss)
    else:
        obj = Tutor.objects.filter(LOGIN=request.session['lid']).update(firstname=firstname, lastname=lastname,
                                                                        DOB=DOB, gender=gender,
                                                                        email=email, phoneno=phoneno, houseno=house,
                                                                        place=place, district=district, state=state,
                                                                        PIN=pin,
                                                                        country=country, qualification=qualification,)
    return HttpResponse('''<script>alert("updated");window.location='/myapp/viewprofile_tutor/'</script>''')

def viewreply(request):
    de=Complaint.objects.filter(FROM_id=request.session['lid'])
    return render(request, "tutor/viewreply.html",{'data':de})
def viewreply_post(request):
    fromdate = request.POST["textfield"]
    todate = request.POST["textfield2"]
    AS = Complaint.objects.filter(date__range=[fromdate, todate])
    return render(request, "admin/viewreply.html", {'data': AS})


def viewstudent_tutor(request):
    return render(request, "tutor/viewstudent.html")
def viewstudent_tutor_post(request):
    return HttpResponse("viewed student")

def viewnote_tutor(request):

    de = Notes.objects.filter(TUTOR__LOGIN_id=request.session['lid'])
    # de=Notes.objects.get(id=id)
    return render(request, "tutor/viewnote.html", {'data': de})
def viewnote_tutor_post(request):
    search = request.POST["textfield"]
    AS = Notes.objects.filter(topic__icontains=search, TUTOR__LOGIN_id=request.session['lid'])
    return render(request, "tutor/viewnote.html", {'data': AS})
def edit_note_tutor(request,id):
    AS = Notes.objects.get(id=id)
    return render(request, "tutor/editnote.html",{'data':AS})
def edit_note_tutor_post(request):
    id = request.POST["id"]
    topic = request.POST["textfield2"]
    description = request.POST["textfield3"]
    if 'fileField' in request.FILES:
        note = request.FILES["fileField2"]
        if note !="":
            from datetime import datetime
            date = datetime.now().strftime('%Y%m%d-%H%M%S') + ".pdf"
            fs = FileSystemStorage()
            fn = fs.save(date, note)
            path = fs.url(date)
            obj = Notes.objects.filter(id=id).update(topic=topic,description=description,note=path,date=date)
    else:
        obj = Notes.objects.filter(id=id).update(topic=topic,description=description)
    de = Notes.objects.filter(TUTOR__LOGIN_id=request.session['lid'])
    return render(request, "tutor/viewnote.html", {'data': de})
def delete_note_tutor(request,id):
    res = Notes.objects.filter(id=id).delete()
    return HttpResponse('''<script>alert("Delete Successfully");window.location='/myapp/viewnote/'</script>''')
def viewdoubts_tutor(request):
    de = Doubts.objects.filter(NOTE__TUTOR__LOGIN_id=request.session['lid'])
    return render(request, "tutor/viewdoubts.html", {'data': de})

def replydoubt(request,id):
    return render(request, "tutor/replydoubts.html", {'data': id})
def replydoubt_post(request):
    reply = request.POST["textfield"]
    id = request.POST["rid"]
    cc = Doubts.objects.filter(id=id).update(reply=reply, status="replied")
    return HttpResponse('''<script>alert("replied");window.location='/myapp/viewdoubts_tutor/'</script>''')

def student_change_password(request):
    return render(request, "student/changepassword.html")
def student_change_password_post(request):
    oldpassword = request.POST["textfield"]
    newpassword = request.POST["textfield2"]
    confirmpassword = request.POST["textfield3"]
    obj =  Login.objects.filter(id = request.session['lid'],password=oldpassword)
    if obj.exists():
        if newpassword == confirmpassword:
            Login.objects.filter(id=request.session['lid']).update(password=newpassword)
            return HttpResponse('''<script>alert("Password Changed");window.location='/myapp/change_password/'</script>''')
        else:
            return HttpResponse('''<script>alert("Password must be same");window.location='/myapp/change_password/'</script>''')
    else:
        return HttpResponse('''<script>alert("Old Password is incorrect");window.location='/myapp/change_password/'</script>''')




def view_tutor_notes(request):
    de = Notes.objects.filter(TUTOR__LOGIN_id=request.session['lid'])
    # de=Notes.objects.get(id=id)
    return render(request, "tutor/viewnote.html", {'data': de})


def inputdoc(request):
    return render(request, "student/inputdoc.html",{'de': ""})


def inputdoc_post(request):
    # document = request.FILES["filefield"]
    # dnm = str(document.name).split(".")[-1]
    # fs = FileSystemStorage()
    # if dnm == "pdf":
    #     import PyPDF2
    #     dt = 'uploads/'+datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')+'pdf'
    #     fs.save(dt,document)
    #
    #     pdf_file = open("/easylearn/media/"+dt, 'rb')
    #     pdf_reader = PyPDF2.PdfReader(pdf_file)
    #     num_pages =  len(pdf_reader.pages)
    #     tt = ""
    #     for i in range(num_pages):
    #         page_obj = pdf_reader.pages[i]
    #         text = page_obj.extract_text()
    #         tt+=text
    #     return render(request,'student/inputdoc.html',{'de':str(tt)})
    # if dnm == 'docx':
    #     import docx
    #     dt = 'uploads/'+datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')+'docx'
    #     fs.save(dt,document)
    #     doc = docx.Document(r"C:\Users\kpana\PycharmProjects\easylearn\media\\"+dt)
    #     fullText = []
    #     for para in doc.paragraphs:
    #         fullText.append(para.text)
    #     return render(request,'student/inputdoc.html',{'de':'\n'.join(fullText)})
    # return HttpResponse('hello')
    document = request.FILES["filefield"]
    file_extension = str(document.name).split(".")[-1]
    fs = FileSystemStorage()

    if file_extension == "pdf":
        dt = 'uploads/' + datetime.datetime.now().strftime('%Y%m%d%H%M%S%f') + 'pdf'
        fs.save(dt, document)

        pdf_file_path = fs.path(dt)
        with open(pdf_file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            text_list = []
            for i in range(num_pages):
                page_obj = pdf_reader.pages[i]
                text_list.append(page_obj.extract_text())

        return render(request, 'student/inputdoc.html', {'de': '\n'.join(text_list)})

    elif file_extension == 'docx':
        dt = 'uploads/' + datetime.datetime.now().strftime('%Y%m%d%H%M%S%f') + 'docx'
        fs.save(dt, document)

        docx_file_path = fs.path(dt)
        doc = docx.Document(docx_file_path)
        text_list = [para.text for para in doc.paragraphs]

        return render(request, 'student/inputdoc.html', {'de': '\n'.join(text_list)})

    return HttpResponse('Invalid file format')


def summary(text):

    res=""
    from transformers import pipeline


    # # using pipeline API for summarization task
    # summarization = pipeline("summarization")
    # original_text = text

    totalstatement = text.split(".")

    l = len(totalstatement) // 25
    k = 0
    while k < l:

        m = ""
        for i in range(k, k + 25):
            m = m + totalstatement[i] + "."

        summarization = pipeline("summarization")
        ln = summarization(m)[0]['summary_text']

        print(ln)
        k = k + 1
        res= res+ln

    return res



def QAgeneration(request):

    text = request.POST["content"]
    from pipelines import pipeline
    nlp = pipeline("question-generation")
    s = nlp(text)

    print(s)

    ss="The generated questions are"

    d=1

    for i in s:

        ss= ss + "Question no "+ str(d) +". Question ."+ i['question'] +". Answer ."+ i['answer']+". "
        d+=1


    d=summary(text)

    ss= ss + ".Generated summary is " + d

    request.session["qa"] = ss

    request.session["status"] = "no"


    request.session["data"]=s
    request.session["d"]= d

    return render(request, "student/QAgenatation.html",{'data': s,'d':d})

def voicreco(request):
    return  render(request,"student/voicereco.html")


def exportwav(request):

    if request.POST["select"] == "wav":
        from gtts import gTTS

        # This module is imported so that we can
        # play the converted audio
        import os
        conversation_history_list = request.session["conversation_history"]

        # Convert the list to a string
        conversation_history_str = "\n".join(
            [f"User: {entry['question']}\nResponse: {entry['response']}\n" for entry in conversation_history_list])
        # The text that you want to convert to audio
        mytext =conversation_history_str

        # Language in which you want to convert
        language = 'en'

        # Passing the text and language to the engine,
        # here we have marked slow=False. Which tells
        # the module that the converted audio should
        # have a high speed
        myobj = gTTS(text=mytext, lang=language, slow=False)
        # Saving the converted audio in a mp3 file named
        # welcome
        myobj.save("\\easylearn\\media\\welcome.mp3")
        request.session["status"] = "ok"
        request.session["file"] = "/media/welcome.mp3"
        conversation_history = request.session["conversation_history"]
        return render(request, "generate new.html", {'conversation_history': conversation_history})
    else:

        from fpdf import FPDF
        s = request.session["conversation_history"]

        # save FPDF() class into a
        # variable pdf
        pdf = FPDF()

        # Add a page
        pdf.add_page()

        # set style and size of font
        # that you want in the pdf
        pdf.set_font("Arial", 'B', size=24)
        pdf.cell(200, 10, txt="Easy Learn", ln=10, align='C')
        pdf.set_font("Arial", 'B', size=15)
        pdf.cell(40, 10, txt="Question & Answers", ln=10, align='L')
        pdf.set_font("Arial", size=12)
        j = 1
        # add another cell
        for i in range(len(s)):
            pdf.multi_cell(0, 10, txt="Q{}: {}".format(j, s[i]['question']), align='L')
            pdf.multi_cell(0, 10, txt="Ans: {}".format(s[i]['answer']), align='L')
            j += 1

        # Split the summary into lines
        summary_lines = request.session["d"].split('.')
        pdf.set_font("Arial", 'B', size=15)
        pdf.cell(40, 10, txt="Summary:", ln=10, align='L')
        # Add each line of the summary to the PDF
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(190, 10, request.session["d"])

        # for line in summary_lines:
        #     pdf.cell(40, 10, txt=line, ln=10, align='L')
        # save the pdf with name .pdf
        pdf.output("C:\\easylearn\\media\\welcome.pdf")
        request.session["status"] = "ok"
        request.session["file"] = "/media/welcome.pdf"

        # from fpdf import FPDF
        #
        # s = request.session["data"]
        #
        # # save FPDF() class into a
        # # variable pdf
        # pdf = FPDF()
        #
        # # Add a page
        # pdf.add_page()
        #
        # # set style and size of font
        # # that you want in the pdf
        # pdf.set_font("Arial", size=15)
        #
        # pdf.cell(200, 10, txt="""Easy Learn""", ln=10, align='C')
        # pdf.cell(40, 10, txt="""Question & Answers""", ln=10, align='L')
        # j = 1
        # # add another cell
        # for i in range(len(s)):
        #     pdf.cell(40, 10, txt="""Q""" + str(j) + " : "+ s[i]['question'], ln=10, align='L')
        #     pdf.cell(40, 10, txt="""Ans :  """+ s[i]['answer'], ln=10, align='L')
        #     j += 1
        # pdf.cell(40, 10, txt="""Summary:""" , ln=10, align='L')
        # pdf.cell(40, 10, txt=request.session["d"], ln=10, align='L')
        # # save the pdf with name .pdf
        # pdf.output("C:\\Users\\kpana\\PycharmProjects\\easylearn\\media\\welcome.pdf")
        #
        # request.session["status"]="ok"
        # request.session["file"]="/media/welcome.pdf"
        return render(request, "student/QAgenatation.html", {'data': request.session["data"], 'd': request.session["d"]})



from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
import os
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import PIL.Image
from django.utils.datastructures import MultiValueDictKeyError



load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))


def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)

    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain


def user_input(request):
    request.session["pdf"] = "no"
    request.session['conversation_history'] = []
    return render(request, "generate new.html",)


def pdf_process(request):
    pdf_docs = request.FILES.getlist('pdf_docs')
    if pdf_docs:
        request.session["pdf"] = "yes"
        text = ""
        for pdf in pdf_docs:
            pdf_reader = PdfReader(pdf)
            for page in pdf_reader.pages:
                text += page.extract_text()

        # raw_text = get_pdf_text(pdf_docs)

        text_chunks = get_text_chunks(text)
        request.session['chunks'] = text_chunks
        get_vector_store(text_chunks)

    conversation_history = request.session.get('conversation_history', [])
    return render(request, "generate new.html", {'conversation_history': conversation_history})


def user_input_post(request):
    user_question = request.POST.get('user_question', '')
    checkbox_status = request.POST.get('toggle')
    try:
        image = request.FILES['image']
    except MultiValueDictKeyError:
        image = None

    conversation_history = request.session.get('conversation_history', [])
    if checkbox_status == "true":
        if request.session["pdf"] == "no":
            placeholder_entry = {'question': user_question, 'response': 'please add a pdf'}
            if conversation_history is None:
                conversation_history = [placeholder_entry]
            else:
                conversation_history.append(placeholder_entry)
        elif user_question == "summarize":
            chunks = request.session.get('chunks')
            summaries = ""
            model = genai.GenerativeModel('gemini-pro')
            for chunk in chunks:
                response = model.generate_content(f"provide a detailed summary of the text. text: {chunk}")
                summary = markdown.markdown(response.text)
                print(summary)
                summaries += summary + '\n'
            conversation_history.append({'question': user_question, 'response': summaries})
            request.session['conversation_history'] = conversation_history

        else:
            embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
            new_db = FAISS.load_local("faiss_index", embeddings)
            docs = new_db.similarity_search(user_question)



            chain = get_conversational_chain()

            response = chain(
                {"input_documents": docs, "question": user_question},
                return_only_outputs=True
            )
            reply_text = response.get('output_text', '')
            reply = markdown.markdown(reply_text)

            conversation_history.append({'question': user_question, 'response': reply})
            request.session['conversation_history'] = conversation_history

    elif image:
        img = PIL.Image.open(image)
        model = genai.GenerativeModel('gemini-pro-vision')
        response = model.generate_content([user_question, img])
        response_text = response.text
        reply = markdown.markdown(response_text)
        conversation_history.append({'question': user_question, 'response': reply})
        request.session['conversation_history'] = conversation_history

    else:
        model = genai.GenerativeModel('gemini-pro')
        response = model.generate_content(user_question)
        reply = markdown.markdown(response.text)
        conversation_history.append({'question': user_question, 'response': reply})
        request.session['conversation_history'] = conversation_history

    return render(request, "generate new.html", {'conversation_history': conversation_history})


def download_conversation(request):
    file_format = request.POST.get('select')
    conversation_history = request.session.get('conversation_history', [])
    cleaned_conversation = [{'question': strip_tags(entry['question']), 'response': strip_tags(entry['response'])} for
                            entry in conversation_history]
    if file_format == 'docx':
        document = Document()
        for entry in cleaned_conversation:
            document.add_paragraph(f"User: {entry['question']}")
            document.add_paragraph(f"Response: {entry['response']}\n")

        # Create a BytesIO buffer to save the document
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Set the appropriate response headers for downloading
        response = HttpResponse(buffer.read(),
                                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=easylearn.docx'
        return response

    elif file_format == 'pdf':
        # Create a BytesIO buffer to save the PDF
        buffer = BytesIO()

        # Create the PDF object using FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        # Set initial y-coordinate
        y = 10
        pdf.set_font("Arial", 'B', size=24)
        pdf.cell(200, 10, txt="Easy Learn", align='C')
        pdf.ln(15)
        pdf.set_font("Arial", size=12)
        # Write the conversation history to the PDF
        for entry in conversation_history:
            # Use MultiCell to handle word wrapping in the user's question
            user_question = f"**User:** {entry['question']}"
            pdf.multi_cell(180, 10, txt=user_question, align='L', markdown=True,  ln=True)

            pdf.ln(2)  # Add some space after the question

            # Set x and y coordinates before rendering the response
            pdf.set_xy(10, pdf.get_y())
            response_text = f"<b>Response:</b>{entry['response']}"
            pdf.write_html(response_text)
            # pdf.multi_cell(0, 10, txt=response_text, align='L', markdown=True, ln=True)

            pdf.ln(10)

        # Output the PDF to the BytesIO buffer
        pdf.output(buffer)

        buffer.seek(0)

        # Set the appropriate response headers for downloading
        response = HttpResponse(buffer.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=easylearn.pdf'
        return response
